"""DOCX segment extract / style-preserving text replace."""

from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from typing import Iterator, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

SKIP_RE = re.compile(
    r"("
    r"^[\d\s\-–—./]+$"  # dates / numbers only
    r"|@"  # email-ish
    r"|https?://"
    r"|^(电话|手机|邮箱|微信|地址|出生|籍贯)[:：]"
    r")"
)


@dataclass
class Segment:
    id: str
    text: str
    kind: str  # paragraph | cell
    editable: bool
    source: str  # docx | pdf


def _iter_block_items(parent) -> Iterator[Tuple[str, object]]:
    """Yield (kind, block) for body paragraphs and tables in document order."""
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, parent)


def _is_editable(text: str) -> bool:
    t = text.strip()
    if len(t) < 8:
        return False
    if SKIP_RE.search(t):
        return False
    # mostly CJK / letters content
    return True


def extract_docx_segments(path: str) -> List[Segment]:
    doc = Document(path)
    segments: List[Segment] = []
    p_idx = 0
    t_idx = 0

    for kind, block in _iter_block_items(doc):
        if kind == "paragraph":
            text = block.text.strip()
            if not text:
                continue
            sid = f"p-{p_idx}"
            segments.append(
                Segment(
                    id=sid,
                    text=text,
                    kind="paragraph",
                    editable=_is_editable(text),
                    source="docx",
                )
            )
            p_idx += 1
        else:
            table: Table = block
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue
                    sid = f"t{t_idx}-r{ri}c{ci}"
                    segments.append(
                        Segment(
                            id=sid,
                            text=text,
                            kind="cell",
                            editable=_is_editable(text),
                            source="docx",
                        )
                    )
            t_idx += 1
    return segments


def _set_paragraph_text_keep_runs(paragraph: Paragraph, new_text: str) -> None:
    """Replace paragraph text while keeping the first run's formatting."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def apply_docx_replacements(src_path: str, dst_path: str, mapping: dict) -> int:
    """Apply id→rewritten map. Returns number of replacements applied."""
    doc = Document(src_path)
    applied = 0
    p_idx = 0
    t_idx = 0
    for kind, block in _iter_block_items(doc):
        if kind == "paragraph":
            text = block.text.strip()
            if not text:
                continue
            sid = f"p-{p_idx}"
            if sid in mapping and mapping[sid] != text:
                _set_paragraph_text_keep_runs(block, mapping[sid])
                applied += 1
            p_idx += 1
        else:
            table: Table = block
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue
                    sid = f"t{t_idx}-r{ri}c{ci}"
                    if sid in mapping and mapping[sid] != text:
                        if cell.paragraphs:
                            _set_paragraph_text_keep_runs(cell.paragraphs[0], mapping[sid])
                            for para in cell.paragraphs[1:]:
                                _set_paragraph_text_keep_runs(para, "")
                        applied += 1
            t_idx += 1
    doc.save(dst_path)
    return applied


def segments_to_dicts(segments: List[Segment]) -> list:
    return [asdict(s) for s in segments]
